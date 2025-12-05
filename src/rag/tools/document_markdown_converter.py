"""
Universal Document to Markdown Converter

Converts any document format (TXT, PDF, Excel, Word, CSV, Database Tables)
to markdown for optimal ChromaDB storage and retrieval using docling-parse.

Benefits:
- Markdown is well-suited for semantic chunking
- Preserves structure and hierarchy with better accuracy
- Better for vector embeddings
- Consistent format across sources
- Supports metadata headers
- docling-parse: AI-powered document parsing for professional quality

Supported Formats:
- PDF: High-fidelity conversion with table/layout preservation
- Excel (XLSX): Multi-sheet support with formatting
- Word (DOCX): Structure and hierarchy preservation
- PowerPoint (PPTX): Slide content extraction
- CSV/TSV: Table formatting
- Text (TXT): Plain text with structure detection
- Database Tables: Structured data with metadata
"""

import json
import pandas as pd
from typing import Dict, List, Any, Optional
from pathlib import Path
from langchain_core.tools import tool
from datetime import datetime

try:
    from docling.document_converter import DocumentConverter
    from docling.datamodel.base_models import Document as DoclingDocument
    HAS_DOCLING = True
except ImportError:
    HAS_DOCLING = False

try:
    from docx import Document as DocxDocument
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False


# ============================================================================
# MARKDOWN CONVERSION UTILITIES
# ============================================================================

class DocumentToMarkdownConverter:
    """Convert various document formats to markdown."""
    
    @staticmethod
    def text_to_markdown(text: str, title: str = "") -> str:
        """
        Convert plain text to markdown.
        
        Args:
            text: Plain text content
            title: Optional title for document
        
        Returns:
            Markdown formatted text
        """
        markdown = ""
        
        if title:
            markdown += f"# {title}\n\n"
        
        # Add metadata header
        markdown += f"_Document converted to markdown on {datetime.now().isoformat()}_\n\n"
        
        # Preserve paragraphs
        paragraphs = text.split('\n\n')
        for para in paragraphs:
            if para.strip():
                # Check if paragraph looks like a heading
                if para.startswith(('##', '###', '#')):
                    markdown += para + "\n\n"
                else:
                    markdown += para + "\n\n"
        
        return markdown
    
    @staticmethod
    def csv_to_markdown(csv_path: str, title: str = "") -> str:
        """
        Convert CSV file to markdown table.
        
        Args:
            csv_path: Path to CSV file
            title: Optional title for document
        
        Returns:
            Markdown formatted text
        """
        try:
            df = pd.read_csv(csv_path)
            
            markdown = ""
            if title:
                markdown += f"# {title}\n\n"
            
            markdown += f"_CSV file converted to markdown table on {datetime.now().isoformat()}_\n\n"
            markdown += f"**Total Records**: {len(df)}\n\n"
            
            # Convert to markdown table
            markdown += df.to_markdown(index=False)
            markdown += "\n\n"
            
            # Add data summary
            markdown += "## Data Summary\n\n"
            for col in df.columns:
                if df[col].dtype in ['int64', 'float64']:
                    markdown += f"- **{col}**: min={df[col].min()}, max={df[col].max()}, mean={df[col].mean():.2f}\n"
                else:
                    unique_count = df[col].nunique()
                    markdown += f"- **{col}**: {unique_count} unique values\n"
            
            return markdown
        except Exception as e:
            return f"# Error Converting CSV\n\nFailed to convert CSV: {str(e)}"
    
    @staticmethod
    def excel_to_markdown(excel_path: str, sheet_name: str = None, title: str = "") -> str:
        """
        Convert Excel file to markdown.
        
        Args:
            excel_path: Path to Excel file
            sheet_name: Specific sheet to convert (None = all sheets)
            title: Optional title for document
        
        Returns:
            Markdown formatted text
        """
        try:
            xls = pd.ExcelFile(excel_path)
            sheets = [sheet_name] if sheet_name else xls.sheet_names
            
            markdown = ""
            if title:
                markdown += f"# {title}\n\n"
            
            markdown += f"_Excel file converted to markdown on {datetime.now().isoformat()}_\n\n"
            markdown += f"**Total Sheets**: {len(xls.sheet_names)}\n\n"
            
            for sheet in sheets:
                df = pd.read_excel(excel_path, sheet_name=sheet)
                
                markdown += f"## Sheet: {sheet}\n\n"
                markdown += f"**Rows**: {len(df)}, **Columns**: {len(df.columns)}\n\n"
                markdown += df.to_markdown(index=False)
                markdown += "\n\n"
            
            return markdown
        except Exception as e:
            return f"# Error Converting Excel\n\nFailed to convert Excel: {str(e)}"
    
    @staticmethod
    def pdf_to_markdown(pdf_path: str, title: str = "") -> str:
        """
        Convert PDF to markdown using docling-parse.
        
        Supports:
        - Text extraction with structure preservation
        - Table extraction and formatting
        - Document layout understanding
        - Multi-page documents
        
        Args:
            pdf_path: Path to PDF file
            title: Optional title for document
        
        Returns:
            Markdown formatted text
        """
        if not HAS_DOCLING:
            return "# Error: docling not installed\n\nInstall with: pip install docling"
        
        try:
            converter = DocumentConverter()
            doc_result = converter.convert_document_or_pdf_path(pdf_path)
            doc: DoclingDocument = doc_result.document
            
            markdown = ""
            if title:
                markdown += f"# {title}\n\n"
            
            markdown += f"_PDF with {len(doc.pages)} pages converted to markdown using docling on {datetime.now().isoformat()}_\n\n"
            
            # Extract markdown directly from docling document
            markdown += doc.export_to_markdown()
            
            return markdown
        except Exception as e:
            return f"# Error Converting PDF\n\nFailed to convert PDF: {str(e)}"
    
    @staticmethod
    def word_to_markdown(docx_path: str, title: str = "") -> str:
        """
        Convert Word document to markdown using docling-parse for better accuracy.
        
        Args:
            docx_path: Path to Word document
            title: Optional title for document
        
        Returns:
            Markdown formatted text
        """
        if not HAS_DOCLING:
            return "# Error: docling not installed\n\nInstall with: pip install docling"
        
        try:
            converter = DocumentConverter()
            doc_result = converter.convert_document_or_pdf_path(docx_path)
            doc: DoclingDocument = doc_result.document
            
            markdown = ""
            if title:
                markdown += f"# {title}\n\n"
            
            markdown += f"_Word document converted to markdown using docling on {datetime.now().isoformat()}_\n\n"
            markdown += doc.export_to_markdown()
            
            return markdown
        except Exception as e:
            # Fallback to python-docx if docling fails
            if HAS_DOCX:
                try:
                    doc = DocxDocument(docx_path)
                    
                    markdown = ""
                    if title:
                        markdown += f"# {title}\n\n"
                    
                    markdown += f"_Word document converted to markdown on {datetime.now().isoformat()}_\n\n"
                    
                    for element in doc.element.body:
                        if element.tag.endswith('p'):
                            para = next(p for p in doc.paragraphs if p._element == element)
                            text = para.text.strip()
                            if text:
                                if para.style.name.startswith('Heading'):
                                    level = int(para.style.name.split()[-1]) if len(para.style.name.split()) > 1 else 2
                                    markdown += f"{'#' * level} {text}\n\n"
                                else:
                                    markdown += f"{text}\n\n"
                        
                        elif element.tag.endswith('tbl'):
                            tbl = next(t for t in doc.tables if t._element == element)
                            rows = []
                            for row in tbl.rows:
                                rows.append([cell.text for cell in row.cells])
                            
                            if rows:
                                df = pd.DataFrame(rows[1:], columns=rows[0])
                                markdown += df.to_markdown(index=False)
                                markdown += "\n\n"
                    
                    return markdown
                except Exception as fallback_err:
                    return f"# Error Converting Word\n\nDocling failed: {str(e)}\n\nFallback failed: {str(fallback_err)}"
            else:
                return f"# Error Converting Word\n\nFailed to convert Word document: {str(e)}"
    
    @staticmethod
    def generic_document_to_markdown(file_path: str, title: str = "") -> str:
        """
        Convert any document format using docling-parse.
        
        Supports: PDF, DOCX, PPTX, HTML, XLSX, and more formats that docling supports.
        
        Args:
            file_path: Path to document file
            title: Optional title for document
        
        Returns:
            Markdown formatted text
        """
        if not HAS_DOCLING:
            return "# Error: docling not installed\n\nInstall with: pip install docling"
        
        try:
            converter = DocumentConverter()
            doc_result = converter.convert_document_or_pdf_path(file_path)
            doc: DoclingDocument = doc_result.document
            
            markdown = ""
            if title:
                markdown += f"# {title}\n\n"
            
            file_type = Path(file_path).suffix.lower()
            markdown += f"_{file_type} document converted to markdown using docling on {datetime.now().isoformat()}_\n\n"
            markdown += doc.export_to_markdown()
            
            return markdown
        except Exception as e:
            return f"# Error Converting Document\n\nFailed to convert {Path(file_path).suffix} document: {str(e)}"
    
    @staticmethod
    def file_to_markdown(file_path: str, title: str = "") -> str:
        """
        Universal method: automatically detect format and convert to markdown.
        
        Uses docling-parse for supported formats, with fallbacks for others.
        
        Args:
            file_path: Path to document file
            title: Optional title for document
        
        Returns:
            Markdown formatted text
        """
        file_ext = Path(file_path).suffix.lower()
        
        if file_ext == ".txt":
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()
            return DocumentToMarkdownConverter.text_to_markdown(content, title)
        
        elif file_ext == ".csv":
            return DocumentToMarkdownConverter.csv_to_markdown(file_path, title)
        
        elif file_ext in [".xlsx", ".xls"]:
            return DocumentToMarkdownConverter.excel_to_markdown(file_path, title=title)
        
        else:
            # Use docling for PDF, DOCX, PPTX, HTML, and others
            return DocumentToMarkdownConverter.generic_document_to_markdown(file_path, title)

    @staticmethod
    def database_table_to_markdown(
        table_data: List[Dict[str, Any]],
        table_name: str,
        column_descriptions: Dict[str, str] = None
    ) -> str:
        """
        Convert database table data to markdown.
        
        Args:
            table_data: List of dictionaries representing table rows
            table_name: Name of the table
            column_descriptions: Optional descriptions for columns
        
        Returns:
            Markdown formatted text
        """
        if not table_data:
            return f"# {table_name}\n\n_Empty table_\n\n"
        
        markdown = f"# {table_name}\n\n"
        markdown += f"_Database table converted to markdown on {datetime.now().isoformat()}_\n\n"
        markdown += f"**Total Records**: {len(table_data)}\n\n"
        
        # Convert to DataFrame for markdown table generation
        df = pd.DataFrame(table_data)
        
        # Add column descriptions if provided
        if column_descriptions:
            markdown += "## Column Descriptions\n\n"
            for col, desc in column_descriptions.items():
                if col in df.columns:
                    markdown += f"- **{col}**: {desc}\n"
            markdown += "\n"
        
        # Add table
        markdown += "## Data\n\n"
        markdown += df.to_markdown(index=False)
        markdown += "\n\n"
        
        # Add statistics for numeric columns
        markdown += "## Statistics\n\n"
        numeric_cols = df.select_dtypes(include=['int64', 'float64']).columns
        if len(numeric_cols) > 0:
            markdown += df[numeric_cols].describe().to_markdown()
        
        return markdown


# ============================================================================
# ENHANCED INGESTION TOOLS WITH MARKDOWN CONVERSION
# ============================================================================

@tool
def convert_to_markdown_tool(
    content: str = None,
    source_type: str = None,
    title: str = "",
    source_path: str = None,
    table_data: List[Dict[str, Any]] = None,
    column_descriptions: Dict[str, str] = None,
    auto_detect: bool = False
) -> str:
    """
    Universal converter to markdown format using docling-parse.
    
    Supports: txt, csv, xlsx, pdf, docx, pptx, html, database_table
    
    Uses docling-parse for professional-grade document conversion with:
    - Layout preservation
    - Table extraction and formatting
    - Multi-page support
    - Hierarchical structure maintenance
    
    Args:
        content: Text content (for text-based sources)
        source_type: Type of source (txt, csv, xlsx, pdf, docx, pptx, html, database_table, auto)
        title: Document title
        source_path: Path to file (for file-based conversions)
        table_data: List of dicts for database table conversion
        column_descriptions: Column metadata for database tables
        auto_detect: If True, auto-detect format from source_path
    
    Returns:
        JSON string with conversion result
    """
    try:
        converter = DocumentToMarkdownConverter()
        markdown = ""
        
        # Auto-detect format if source_path provided
        if auto_detect and source_path:
            markdown = converter.file_to_markdown(source_path, title)
        
        elif source_type.lower() == "txt":
            markdown = converter.text_to_markdown(content or "", title)
        
        elif source_type.lower() == "csv":
            if not source_path:
                # Try to parse CSV from content string
                import io
                df = pd.read_csv(io.StringIO(content))
                markdown = f"# {title}\n\n" if title else ""
                markdown += df.to_markdown(index=False)
            else:
                markdown = converter.csv_to_markdown(source_path, title)
        
        elif source_type.lower() in ["xlsx", "excel"]:
            if source_path:
                markdown = converter.excel_to_markdown(source_path, title=title)
            else:
                return json.dumps({
                    "success": False,
                    "error": "Excel conversion requires source_path"
                })
        
        elif source_type.lower() in ["pdf", "docx", "pptx", "html", "word"]:
            if source_path:
                # Use generic docling converter for all these formats
                markdown = converter.generic_document_to_markdown(source_path, title)
            else:
                return json.dumps({
                    "success": False,
                    "error": f"{source_type} conversion requires source_path"
                })
        
        elif source_type.lower() == "database_table":
            if table_data:
                markdown = converter.database_table_to_markdown(
                    table_data,
                    title,
                    column_descriptions
                )
            else:
                return json.dumps({
                    "success": False,
                    "error": "Database table conversion requires table_data"
                })
        
        else:
            return json.dumps({
                "success": False,
                "error": f"Unsupported source type: {source_type}"
            })
        
        return json.dumps({
            "success": True,
            "markdown": markdown,
            "source_type": source_type,
            "length_chars": len(markdown),
            "message": f"Successfully converted {source_type} to markdown using docling"
        })
    
    except Exception as e:
        return json.dumps({
            "success": False,
            "error": str(e),
            "source_type": source_type
        })


if __name__ == "__main__":
    # Example usage
    converter = DocumentToMarkdownConverter()
    
    # Convert sample text
    sample_text = "This is a sample document.\n\nIt has multiple paragraphs."
    markdown = converter.text_to_markdown(sample_text, "Sample Document")
    print("Text Conversion:")
    print(markdown)
    print("\n" + "="*80 + "\n")
    
    # Convert sample database table
    sample_data = [
        {"id": 1, "name": "John", "dept": "Engineering"},
        {"id": 2, "name": "Jane", "dept": "Finance"},
        {"id": 3, "name": "Bob", "dept": "HR"}
    ]
    descriptions = {
        "id": "Employee ID",
        "name": "Employee Name",
        "dept": "Department"
    }
    markdown = converter.database_table_to_markdown(sample_data, "Employees", descriptions)
    print("Database Table Conversion:")
    print(markdown)
